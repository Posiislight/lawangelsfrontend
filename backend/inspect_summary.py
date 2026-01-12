from docx import Document
import os

base_dir = os.path.dirname(os.path.abspath(__file__))
summary_dir = os.path.join(base_dir, 'summarynotes')
files = [f for f in os.listdir(summary_dir) if f.endswith('.docx')]
if not files:
    print("No docx found")
    exit(1)
    
filepath = os.path.join(summary_dir, files[0])

doc = Document(filepath)

print(f"Total Paragraphs: {len(doc.paragraphs)}")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if style.startswith('Heading 1'):
        print(f"P{i} [{style}]: {text}")
