"""
Management command to import summary notes from .docx files.
Parses Word documents and creates SummaryNotes and SummaryNotesChapter records.
"""
import os
import re
from django.core.management.base import BaseCommand
from django.conf import settings
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from quiz.summary_notes_models import SummaryNotes, SummaryNotesChapter


# Mapping of docx filenames to subject info
SUBJECT_MAPPING = {
    'BUSINESS LAW SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Business Law',
        'subject': 'Business Law',
        'category': 'FLK1',
        'icon': '💼',
        'order': 1,
    },
    # Filename with double space
    'BUSINESS LAW  SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Business Law',
        'subject': 'Business Law',
        'category': 'FLK1',
        'icon': '💼',
        'order': 1,
    },
    'CONSTITUTIONAL AND ADMINISTRATIVE LAW SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Constitutional & Administrative Law',
        'subject': 'Constitutional Law',
        'category': 'FLK1',
        'icon': '⚖️',
        'order': 2,
    },
    'CRIMINAL LAW SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Criminal Law',
        'subject': 'Criminal Law',
        'category': 'FLK1',
        'icon': '🔒',
        'order': 3,
    },
    'CRIMINAL PRACTICE SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Criminal Practice',
        'subject': 'Criminal Practice',
        'category': 'FLK2',
        'icon': '👨‍⚖️',
        'order': 4,
    },
    'DISPUTE RESOLUTION SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Dispute Resolution',
        'subject': 'Dispute Resolution',
        'category': 'FLK2',
        'icon': '🤝',
        'order': 5,
    },
    'LAND LAW SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Land Law',
        'subject': 'Land Law',
        'category': 'FLK1',
        'icon': '🏠',
        'order': 6,
    },
    'LAW ANGELS- SOLICITORS\' ACCOUNT.docx': {
        'title': 'Solicitors\' Accounts',
        'subject': 'Solicitors Accounts',
        'category': 'FLK2',
        'icon': '📊',
        'order': 7,
    },
    'LEGAL SERVICES SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Legal Services',
        'subject': 'Legal Services',
        'category': 'FLK1',
        'icon': '📋',
        'order': 8,
    },
    'LEGAL SYSTEM OF ENGLAND AND WALES SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Legal System of England & Wales',
        'subject': 'Legal System',
        'category': 'FLK1',
        'icon': '🏛️',
        'order': 9,
    },
    'PROFESSIONAL ETHICS SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Professional Ethics',
        'subject': 'Ethics',
        'category': 'FLK1',
        'icon': '📜',
        'order': 10,
    },
    'PROPERTY PRACTICE SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Property Practice',
        'subject': 'Property Practice',
        'category': 'FLK2',
        'icon': '🏢',
        'order': 11,
    },
    'TORTS SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Torts',
        'subject': 'Torts',
        'category': 'FLK1',
        'icon': '⚠️',
        'order': 12,
    },
    'TRUSTS SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Trusts',
        'subject': 'Trusts',
        'category': 'FLK1',
        'icon': '🔐',
        'order': 13,
    },
    'WILLS AND ADMINISTRATION OF ESTATES SUMMARY WITHOUT BORDERS.docx': {
        'title': 'Wills & Administration of Estates',
        'subject': 'Wills',
        'category': 'FLK2',
        'icon': '📝',
        'order': 14,
    },
}


class Command(BaseCommand):
    help = 'Import summary notes from .docx files in the summarynotes directory'

    def add_arguments(self, parser):
        parser.add_argument(
            '--clear',
            action='store_true',
            help='Clear existing summary notes before importing',
        )

    def handle(self, *args, **options):
        summary_notes_dir = os.path.join(settings.BASE_DIR, 'summarynotes')
        
        if not os.path.exists(summary_notes_dir):
            self.stderr.write(self.style.ERROR(f'Directory not found: {summary_notes_dir}'))
            return
        
        if options['clear']:
            self.stdout.write('Clearing existing summary notes...')
            SummaryNotesChapter.objects.all().delete()
            SummaryNotes.objects.all().delete()
        
        # Get all docx files
        docx_files = [f for f in os.listdir(summary_notes_dir) if f.endswith('.docx')]
        self.stdout.write(f'Found {len(docx_files)} .docx files')
        
        for filename in docx_files:
            filepath = os.path.join(summary_notes_dir, filename)
            self.import_docx(filepath, filename)
        
        self.stdout.write(self.style.SUCCESS('Import complete!'))

    def import_docx(self, filepath, filename):
        """Import a single .docx file."""
        self.stdout.write(f'Processing: {filename}')
        
        # Get subject info from mapping
        subject_info = SUBJECT_MAPPING.get(filename)
        if not subject_info:
            self.stdout.write(self.style.WARNING(f'  No mapping found for {filename}, skipping'))
            return
        
        # Create or update SummaryNotes record
        summary_notes, created = SummaryNotes.objects.update_or_create(
            source_file=filename,
            defaults={
                'title': subject_info['title'],
                'subject': subject_info['subject'],
                'category': subject_info['category'],
                'icon': subject_info['icon'],
                'order': subject_info['order'],
                'description': f"Study notes for {subject_info['title']}",
            }
        )
        
        action = 'Created' if created else 'Updated'
        self.stdout.write(f'  {action} summary notes: {summary_notes.title}')
        
        # Parse the document
        try:
            doc = Document(filepath)
        except Exception as e:
            self.stderr.write(self.style.ERROR(f'  Error opening document: {e}'))
            return
        
        # Delete existing chapters
        summary_notes.chapters.all().delete()
        
        # Parse paragraphs into chapters
        chapters = self.parse_chapters(doc)
        
        # Create chapter records
        for idx, chapter in enumerate(chapters):
            SummaryNotesChapter.objects.create(
                summary_notes=summary_notes,
                title=chapter['title'],
                order=idx + 1,
                content=chapter['content'],
            )
        
        self.stdout.write(f'  Created {len(chapters)} chapters')

    def parse_chapters(self, doc):
        """
        Parse a Word document into chapters based on heading styles.
        Uses Heading 1 as chapter breaks.
        """
        chapters = []
        current_chapter = None
        current_content = []
        
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            text = para.text.strip()
            
            # Check if this is a heading (chapter break)
            if style_name.startswith('Heading 1') or (
                style_name == 'Normal' and len(text) < 100 and text.isupper()
            ):
                # Save previous chapter
                if current_chapter and current_content:
                    chapters.append({
                        'title': current_chapter,
                        'content': '\n'.join(current_content)
                    })
                
                # Start new chapter
                current_chapter = text if text else 'Introduction'
                current_content = []
            elif text:
                # Add content to current chapter
                html_content = self.paragraph_to_html(para)
                if html_content:
                    current_content.append(html_content)
        
        # Don't forget the last chapter
        if current_chapter and current_content:
            chapters.append({
                'title': current_chapter,
                'content': '\n'.join(current_content)
            })
        
        # If no chapters found, create a single chapter with all content
        if not chapters:
            all_content = []
            for para in doc.paragraphs:
                html_content = self.paragraph_to_html(para)
                if html_content:
                    all_content.append(html_content)
            
            if all_content:
                chapters.append({
                    'title': 'Summary Notes',
                    'content': '\n'.join(all_content)
                })
        
        return chapters

    def paragraph_to_html(self, para):
        """Convert a paragraph to HTML."""
        text = para.text.strip()
        if not text:
            return ''
        
        style_name = para.style.name if para.style else ''
        
        # Handle different styles
        if style_name.startswith('Heading 2'):
            return f'<h2>{self.escape_html(text)}</h2>'
        elif style_name.startswith('Heading 3'):
            return f'<h3>{self.escape_html(text)}</h3>'
        elif style_name.startswith('Heading'):
            return f'<h4>{self.escape_html(text)}</h4>'
        elif 'List' in style_name or text.startswith('•') or text.startswith('-'):
            # Handle list items
            clean_text = text.lstrip('•-').strip()
            return f'<li>{self.escape_html(clean_text)}</li>'
        else:
            # Regular paragraph - check for bold/italic runs
            html_parts = []
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                
                if run.bold and run.italic:
                    html_parts.append(f'<strong><em>{self.escape_html(run_text)}</em></strong>')
                elif run.bold:
                    html_parts.append(f'<strong>{self.escape_html(run_text)}</strong>')
                elif run.italic:
                    html_parts.append(f'<em>{self.escape_html(run_text)}</em>')
                else:
                    html_parts.append(self.escape_html(run_text))
            
            if html_parts:
                return f'<p>{"".join(html_parts)}</p>'
            return f'<p>{self.escape_html(text)}</p>'

    def escape_html(self, text):
        """Escape HTML special characters."""
        return (
            text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;')
        )
