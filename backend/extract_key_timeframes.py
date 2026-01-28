
import docx
import os
import glob

def extract_text(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    row_text.append(text)
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return '\n'.join(full_text)

def main():
    directory = r"c:\Users\adele\lawangelsfrontend\backend\key deadline and timeframes"
    docx_files = glob.glob(os.path.join(directory, "*.docx"))
    
    with open('extracted_timeframes.txt', 'w', encoding='utf-8') as f:
        for file_path in docx_files:
            print(f"Processing {os.path.basename(file_path)}...")
            try:
                text = extract_text(file_path)
                f.write(f"--- CONTENT OF {os.path.basename(file_path)} ---\n")
                f.write(text + "\n")
                f.write("------------------------------------------------\n")
            except Exception as e:
                print(f"Error reading {file_path}: {e}")

if __name__ == "__main__":
    main()
