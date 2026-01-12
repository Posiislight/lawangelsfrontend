from docx import Document
import os

filepath = os.path.join('backend', 'flashcards', 'flk1', 'BUSINESS LAW FLASHCARDS (1).docx')
doc = Document(filepath)

print(f"Paragraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs[:20]):
    if p.text.strip():
        print(f"P{i}: {repr(p.text)}")

print(f"Tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows")
    for r_idx, row in enumerate(table.rows[:5]):
        row_text = ' | '.join([cell.text.strip() for cell in row.cells])
        print(f"  R{r_idx}: {row_text}")
